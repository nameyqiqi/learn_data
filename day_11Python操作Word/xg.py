from docx import Document
# 从docx中导入颜色的包
from docx.shared import RGBColor
# 从docx中导入修改文本位置的包
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

word = Document()

# 一、修改标题样式
word.add_heading('标题', level=0)
# 修改标题的颜色color --> 计算机中比较热衷于颜色的十六进制或rgb（红绿蓝三原色）
# 三原色分别有265个等级，0~256
# 修改word文档的标题（Title）的样式（styles）中的字体（font）的rgb颜色
word.styles['Title'].font.color.rgb = RGBColor(255, 255, 0)
# 左对齐、右对齐、居中对齐：LEFT、RIGHT、CENTER
word.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 二、修改表格样式
word.add_table(rows=2, cols=2, style='Light Grid')

word.save('./DOCX/样式修改.docx')