# doc,docx --> word文档的后缀名
# Python操作word  --> python-docx模块
# 写代码时，导入python-docx模块需要用它的另外一个名字：docx
from docx import Document
#创建文档
doc_1 = Document()
# 添加标题： add_heading level 0~9 1-9级标题 数字越大等级越小（一级最大），字号越小
doc_1.add_heading('这是标题',level=0)
# 四. 写入正文
# 古诗的题目：也算标题
doc_1.add_heading('望庐山瀑布',level=1)
# 写入内容（段落）
p_1 = doc_1.add_paragraph('日照香炉生紫烟，\n遥看瀑布挂前川。')
#在段落前添加新内容（在python操作word中合适的位置适当的创建变量）
p_1.insert_paragraph_before('唐 李白')
# 续写段落
p_1.add_run('\n飞流直下三千尺，\n疑是银河落九天。')
# 添加分页符
doc_1.add_page_break()
# 添加图片
doc_1.add_picture('./DOCX/1.PNG')
# 图片大小与word大小不一致 修改图片样式
# Cm厘米 Pt磅 Inches英寸 单位 Inches>Cm>Pt
from docx.shared import Cm,Inches,Pt
doc_1.add_picture('./DOCX/1.PNG',width=Pt(400))
doc_1.add_picture('./DOCX/1.PNG',width=Cm(15))
doc_1.add_picture('./DOCX/1.PNG',width=Inches(5.0))

# 添加表格：add_table(rows, cols) -> 创建rows行cols列的一个表格
t1 = doc_1.add_table(rows=1, cols=2)
# 表格.rows[N] --> 获取表格的第N + 1行
r1 = t1.rows[0]
r1.cells[0].text = '作者'
# 行.cells[N] --> 获取行的第N + 1列
# text：指定单元格的值
r1.cells[1].text = '古诗'

# add_row():添加行，添加行的时候自动生成列
data = [
    ['李白', '将进酒'], ['李白', '望月'], ['李白', '蜀道难']
]
for i in data:
    # 添加一行，同时获取这一行的所有单元格
    l = t1.add_row().cells
    # 借助循环，通过下标形式将每个元素添加到每个单元格
    for j in range(0, len(i)):
        l[j].text = i[j]

# 最终、关闭文件
doc_1.save('./DOCX/test.docx')
