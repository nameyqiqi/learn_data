from docx import Document
import  time
t = time.localtime()
edate=datatime = f'({t.tm_year}年{t.tm_mon}月{t.tm_mday}日)'
#print(datatime)

doc_1 = Document()
doc_1.add_heading('离职证明',level=1)
name = ['张三','李四','王五']
for i in name:
    id = input('please input identidy card:')
    sdate = input('please input sdate:')
    department = input('please input department:')
    position = input('please input position:')
    p_1 = doc_1.add_paragraph(f'  兹证明{i}，身份证号码: {id}，于{sdate}至{edate}在我单位'
                          f'{department}部门担任{position}职务，在职期间无不良表现。因个人'
                          f'原因，于{edate}起终止解除劳动合同。现已结清财务相关费用，办'
                          f'理完解除劳动关系相关手续，双方不存在任何劳动争议')
    doc_1.add_paragraph('特此证明!')
    doc_1.add_paragraph(f'公司(盖章):学习梦想\n{edate}')

    doc_1.save(f'./DOCX/{i}离职证明.docx')